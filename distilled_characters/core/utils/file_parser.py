"""File parser — extracts text from PDF, DOCX, TXT, and other formats."""
from __future__ import annotations

import io
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_file(file_path: str | Path) -> str:
    """Parse a file and return its text content. Supports PDF, TXT, MD, DOCX, DOC.

    Returns empty string if parsing fails or format is unsupported.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md", ".markdown", ".csv", ".json", ".xml", ".html", ".htm", ".py", ".js", ".ts", ".rst"):
        return _parse_text(path)

    if suffix == ".pdf":
        return _parse_pdf(path)

    if suffix == ".docx":
        return _parse_docx(path)

    if suffix == ".doc":
        return _parse_doc(path)

    if suffix in (".xmind", ".xmnd"):
        return _parse_xmind_bytes(path.read_bytes())

    # Fallback: try as text
    return _parse_text(path)


def parse_bytesio(filename: str, data: bytes) -> str:
    """Parse file content from bytes (in-memory upload).

    Supported formats: PDF, DOCX, TXT, MD, CSV, JSON, HTML, and more.
    Raises ValueError for unsupported/unparseable files.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        text = _parse_pdf_bytes(data)
        if not text.strip():
            raise ValueError("PDF 无法解析 — 文件可能是扫描图片（已尝试 OCR，最多 20 页）、加密或已损坏。请尝试复制其中文本直接粘贴。")
        return text

    if suffix in (".docx", ".doc"):
        text = _parse_docx_bytes(data)
        if not text.strip():
            raise ValueError("DOCX/DOC 无法解析 — 文件可能为空或格式异常。")
        return text

    if suffix in (".xmind", ".xmnd"):
        text = _parse_xmind_bytes(data)
        if not text.strip():
            raise ValueError("XMind 文件无法解析 — 可能是空脑图或不支持的版本（仅支持 XMind 8 / Zen / 2020+）。")
        return text

    # Default: decode as text
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = data.decode("gbk")
        except UnicodeDecodeError:
            try:
                text = data.decode("gb2312")
            except UnicodeDecodeError:
                text = data.decode("utf-8", errors="replace")

    if not text.strip():
        raise ValueError(f"文件为空或无法识别编码")
    return text


# ── Text ──────────────────────────────────────────────────────────

def _parse_text(path: Path) -> str:
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


# ── PDF ───────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> str:
    try:
        return _parse_pdf_bytes(path.read_bytes())
    except Exception:
        return ""


def _parse_pdf_bytes(data: bytes) -> str:
    """Extract text from PDF bytes. Tries multiple backends in order:

    1. pdfplumber — best for Chinese text and complex layouts (extracts tables too)
    2. pypdf — lightweight, good fallback
    3. OCR via pdf2image + pytesseract — for scanned/image-based PDFs
    """
    # 1. pdfplumber (preferred: handles Chinese, tables, complex layouts)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            parts = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    parts.append(text)
                try:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = "\n".join(
                                " | ".join(cell or "" for cell in row)
                                for row in table
                            )
                            parts.append(table_text)
                except Exception:
                    pass
            result = "\n\n".join(parts)
            if result.strip():
                logger.info("PDF parsed via pdfplumber: %d pages, %d chars", len(pdf.pages), len(result))
                return result
    except ImportError:
        logger.debug("pdfplumber not installed")
    except Exception as e:
        logger.warning("pdfplumber failed: %s", e)

    # 2. pypdf (good for text-based PDFs)
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        result = "\n\n".join(parts)
        if result.strip():
            logger.info("PDF parsed via pypdf: %d pages, %d chars", len(reader.pages), len(result))
            return result
    except ImportError:
        logger.debug("pypdf not installed")
    except Exception as e:
        logger.warning("pypdf failed: %s", e)

    # 3. OCR — for scanned/image-based PDFs (common: FreePic2Pdf, scanner output)
    text = _pdf_ocr(data)
    if text.strip():
        return text

    return ""


# ── OCR for scanned PDFs ──────────────────────────────────────────

def _pdf_ocr(data: bytes, max_pages: int = 20) -> str:
    """OCR a PDF by rendering pages to images and running Tesseract.

    Only invoked when all text-extraction backends return empty.
    Capped at max_pages to prevent unbounded time on large documents.
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as e:
        logger.debug("OCR deps not installed: %s", e)
        return ""

    try:
        images = convert_from_bytes(data, dpi=200, first_page=1, last_page=max_pages + 1)
    except Exception as e:
        logger.warning("pdf2image conversion failed: %s", e)
        return ""

    parts = []
    for i, img in enumerate(images):
        try:
            # chi_sim+eng handles mixed Chinese/English content
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            if text.strip():
                parts.append(text.strip())
        except Exception as e:
            logger.warning("OCR page %d failed: %s", i + 1, e)

    result = "\n\n".join(parts)
    if result.strip():
        logger.info("PDF OCR'd: %d pages, %d chars", len(images), len(result))
    return result


# ── DOCX ──────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> str:
    try:
        return _parse_docx_bytes(path.read_bytes())
    except Exception:
        return ""


def _parse_docx_bytes(data: bytes) -> str:
    """Extract text from DOCX bytes. Falls back to zipfile raw XML parsing."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))

        # Collect paragraphs
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also collect text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text)
                if row_text.strip():
                    parts.append(row_text)

        result = "\n".join(parts)
        if result.strip():
            logger.info("DOCX parsed via python-docx: %d paragraphs, %d chars", len(parts), len(result))
            return result
    except ImportError:
        logger.debug("python-docx not installed")
    except Exception as e:
        logger.warning("python-docx failed: %s", e)

    # Fallback: DOCX is a ZIP of XML files. Extract text from document.xml.
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            if "word/document.xml" in z.namelist():
                xml_content = z.read("word/document.xml")
                root = ET.fromstring(xml_content)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                parts = []
                for t_elem in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                    if t_elem.text:
                        parts.append(t_elem.text)
                result = "".join(parts)
                if result.strip():
                    logger.info("DOCX parsed via XML fallback: %d chars", len(result))
                    return result
    except Exception as e:
        logger.warning("DOCX XML fallback failed: %s", e)

    return ""


# ── DOC (legacy .doc, not .docx) ──────────────────────────────────

def _parse_doc(path: Path) -> str:
    """Parse legacy .doc files (pre-2007 Word format).

    Requires antiword or textract. On macOS: brew install antiword
    """
    # Try antiword first
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            logger.info("DOC parsed via antiword: %d chars", len(result.stdout))
            return result.stdout
    except FileNotFoundError:
        logger.debug("antiword not installed (brew install antiword)")
    except Exception as e:
        logger.warning("antiword failed: %s", e)

    # Try textract
    try:
        import textract
        text = textract.process(str(path)).decode("utf-8")
        if text.strip():
            logger.info("DOC parsed via textract: %d chars", len(text))
            return text
    except ImportError:
        logger.debug("textract not installed")
    except Exception as e:
        logger.warning("textract failed: %s", e)

    return ""


# ── XMind (.xmind / .xmnd) ─────────────────────────────────────────

def _parse_xmind_bytes(data: bytes) -> str:
    """Extract topic text from an XMind file. XMind files are ZIP archives.

    Two layouts are supported:
    - XMind Zen / 2020+ : ``content.json`` (list of sheets with rootTopic)
    - XMind 8 / legacy  : ``content.xml`` (xmap namespace)

    Returns indented Markdown-like outline with one node per line so the LLM
    keeps the parent/child relationships.
    """
    import io
    import json
    import zipfile

    lines: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            names = z.namelist()
            if "content.json" in names:
                raw = z.read("content.json")
                payload = json.loads(raw.decode("utf-8", errors="replace"))
                # `payload` is a list of sheets
                sheets = payload if isinstance(payload, list) else [payload]
                for sheet in sheets:
                    title = (sheet.get("title") or "").strip()
                    if title:
                        lines.append(f"# {title}")
                    root = sheet.get("rootTopic") or {}
                    _xmind_zen_walk(root, lines, depth=0)
                    lines.append("")
            elif "content.xml" in names:
                raw = z.read("content.xml")
                _xmind_legacy_walk(raw, lines)
            else:
                logger.warning("XMind: neither content.json nor content.xml found in %s", names[:5])
                return ""
    except zipfile.BadZipFile:
        logger.warning("XMind: file is not a valid ZIP archive")
        return ""
    except Exception as e:
        logger.warning("XMind parsing failed: %s", e)
        return ""

    result = "\n".join(lines).strip()
    if result:
        logger.info("XMind parsed: %d nodes, %d chars", len(lines), len(result))
    return result


def _xmind_zen_walk(node: dict, lines: list[str], depth: int) -> None:
    if not isinstance(node, dict):
        return
    title = (node.get("title") or "").strip()
    notes = node.get("notes") or {}
    note_text = ""
    if isinstance(notes, dict):
        plain = notes.get("plain") or {}
        if isinstance(plain, dict):
            note_text = (plain.get("content") or "").strip()
    indent = "  " * depth
    if title:
        lines.append(f"{indent}- {title}")
    if note_text:
        for ln in note_text.splitlines():
            ln = ln.strip()
            if ln:
                lines.append(f"{indent}  > {ln}")

    children = node.get("children") or {}
    if isinstance(children, dict):
        for key in ("attached", "summary", "callout", "detached"):
            for child in children.get(key) or []:
                _xmind_zen_walk(child, lines, depth + 1)


def _xmind_legacy_walk(xml_bytes: bytes, lines: list[str]) -> None:
    from xml.etree import ElementTree as ET

    root = ET.fromstring(xml_bytes)
    ns = {"x": "urn:xmind:xmap:xmlns:content:2.0"}

    def text_of(elem, tag: str) -> str:
        child = elem.find(f"x:{tag}", ns)
        return (child.text or "").strip() if child is not None and child.text else ""

    def walk(topic, depth: int) -> None:
        title = text_of(topic, "title")
        if title:
            lines.append(f"{'  ' * depth}- {title}")
        note = topic.find("x:notes/x:plain", ns)
        if note is not None and note.text:
            for ln in note.text.splitlines():
                ln = ln.strip()
                if ln:
                    lines.append(f"{'  ' * (depth + 1)}> {ln}")
        # Topic children live under <children><topics type="attached"><topic/>...
        for topics_elem in topic.findall("x:children/x:topics", ns):
            for child in topics_elem.findall("x:topic", ns):
                walk(child, depth + 1)

    for sheet in root.findall("x:sheet", ns):
        sheet_title = text_of(sheet, "title")
        if sheet_title:
            lines.append(f"# {sheet_title}")
        for topic in sheet.findall("x:topic", ns):
            walk(topic, 0)
        lines.append("")
